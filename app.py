import streamlit as st
import fitz
import time
from langchain_groq import ChatGroq
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document
import io

# ── Page Config ──────────────────────────────────────────
st.set_page_config(
    page_title="AI Legal Document Simplifier",
    page_icon="⚖️",
    layout="wide"
)

# ── Styling ───────────────────────────────────────────────
st.markdown("""
    <style>
    .main { background-color: #f9f9f9; }
    .stButton>button {
        background-color: #2c3e50;
        color: white;
        border-radius: 8px;
        padding: 0.5em 1.5em;
        font-size: 16px;
    }
    .flag-box {
        background-color: #fff3cd;
        border-left: 5px solid #ffc107;
        padding: 10px 15px;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .summary-box {
        background-color: #e8f4f8;
        border-left: 5px solid #2980b9;
        padding: 15px;
        border-radius: 5px;
    }
    </style>
""", unsafe_allow_html=True)

# ── Helper Functions ──────────────────────────────────────
def extract_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page_num, page in enumerate(doc, start=1):
        text += f"\n--- Page {page_num} ---\n{page.get_text()}"
    doc.close()
    return text

def split_text(text):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", " "]
    )
    return splitter.split_text(text)

def simplify_chunk(chunk, llm):
    messages = [
        SystemMessage(content="""You are a legal document simplifier.
Convert complex legal text into simple plain English anyone can understand.
- Use short sentences
- Avoid legal jargon
- Highlight important obligations, deadlines, or penalties with ⚠️
- Be concise but don't miss key points"""),
        HumanMessage(content=f"Simplify this text:\n\n{chunk}")
    ]
    return llm.invoke(messages).content

def detect_red_flags(chunk, llm):
    messages = [
        SystemMessage(content="""You are a helpful legal assistant reviewing a document on behalf of a regular person.
Extract ANY clause a normal person should be aware of before signing or agreeing.
This includes penalties, auto-renewal, data sharing, termination conditions, obligations, liability limits, deadlines.
Format: 🚩 [Clause Type]: [Simple one line explanation]
Only respond with NONE if the chunk contains absolutely no actionable information."""),
        HumanMessage(content=f"Review this section:\n\n{chunk}")
    ]
    return llm.invoke(messages).content.strip()

def generate_summary(simplified_chunks, llm):
    sampled = (simplified_chunks[:5] +
               simplified_chunks[len(simplified_chunks)//2 - 2 : len(simplified_chunks)//2 + 2] +
               simplified_chunks[-5:])
    combined = "\n\n".join(sampled)
    messages = [
        SystemMessage(content="""You are a professional document analyst.
Write a clear executive summary structured as:
📌 **What this document is about** (2-3 sentences)
📊 **Key highlights** (4-5 bullet points)
📅 **Important figures or dates mentioned** (if any)
🎯 **Who this document is relevant for** (1-2 sentences)"""),
        HumanMessage(content=f"Summarize these sections:\n\n{combined}")
    ]
    return llm.invoke(messages).content

def create_docx(summary, red_flags, simplified_chunks):
    doc = Document()
    doc.add_heading("AI Legal Document Simplifier", level=0)
    doc.add_paragraph("For reference only — not legal advice.")
    doc.add_paragraph("")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_paragraph("")

    if red_flags:
        doc.add_heading("Red Flags Detected", level=1)
        for flag in red_flags:
            doc.add_paragraph(flag)
        doc.add_paragraph("")

    doc.add_heading("Simplified Sections", level=1)
    for i, section in enumerate(simplified_chunks, start=1):
        doc.add_heading(f"Section {i}", level=2)
        doc.add_paragraph(section)
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ── Main App ──────────────────────────────────────────────
st.title("⚖️ AI Legal Document Simplifier")
st.markdown("Upload any legal or formal document — get a plain English summary, red flag alerts, and a downloadable simplified report.")
st.markdown("---")

# Sidebar
with st.sidebar:
    st.header("⚙️ Settings")
    api_key = st.text_input("Groq API Key", type="password", placeholder="gsk_...")
    model = st.selectbox("Model", ["llama-3.1-8b-instant", "llama3-8b-8192", "mixtral-8x7b-32768"])
    st.markdown("---")
    st.markdown("### 📌 How to use")
    st.markdown("1. Enter your Groq API key\n2. Upload a PDF\n3. Click **Analyze**\n4. Download your report")
    st.markdown("---")
    st.caption("⚠️ This tool is for reference only and does not constitute legal advice.")

# File Upload
uploaded_file = st.file_uploader("📄 Upload PDF Document", type=["pdf"])

if uploaded_file and api_key:
    if st.button("🔍 Analyze Document"):

        llm = ChatGroq(api_key=api_key, model_name=model, temperature=0.3)

        # Step 1 — Extract
        with st.spinner("Extracting text from PDF..."):
            raw_text = extract_text(uploaded_file)
            chunks = split_text(raw_text)
        st.success(f"✅ Extracted {len(chunks)} sections from document")

        # Step 2 — Simplify
        st.markdown("### 🔄 Simplifying Document...")
        progress_bar = st.progress(0)
        status_text = st.empty()
        simplified_chunks = []

        for i, chunk in enumerate(chunks):
            simplified_chunks.append(simplify_chunk(chunk, llm))
            progress_bar.progress((i + 1) / len(chunks))
            status_text.text(f"Processing section {i+1} of {len(chunks)}...")
            time.sleep(0.5)

        status_text.text("✅ Simplification complete!")

        # Step 3 — Red Flags
        st.markdown("### 🚩 Scanning for Red Flags...")
        red_flags = []
        flag_bar = st.progress(0)

        for i, chunk in enumerate(chunks):
            result = detect_red_flags(chunk, llm)
            if result != "NONE":
                red_flags.append(f"Section {i+1}: {result}")
            flag_bar.progress((i + 1) / len(chunks))
            time.sleep(0.5)

        # Step 4 — Summary
        with st.spinner("Generating executive summary..."):
            summary = generate_summary(simplified_chunks, llm)

        st.markdown("---")

        # ── Results ──
        col1, col2, col3 = st.columns(3)
        col1.metric("📄 Total Sections", len(chunks))
        col2.metric("🚩 Red Flags Found", len(red_flags))
        col3.metric("📝 Summary Generated", "Yes")

        st.markdown("---")

        # Summary
        st.markdown("### 📊 Executive Summary")
        st.markdown(f'<div class="summary-box">{summary}</div>', unsafe_allow_html=True)

        # Red Flags
        st.markdown("### 🚩 Red Flags")
        if red_flags:
            for flag in red_flags:
                st.markdown(f'<div class="flag-box">{flag}</div>', unsafe_allow_html=True)
        else:
            st.info("✅ No red flags found — this document appears to be informational.")

        # Simplified Sections
        st.markdown("### 📄 Simplified Document")
        for i, section in enumerate(simplified_chunks, start=1):
            with st.expander(f"Section {i}"):
                st.write(section)

        # Download
        st.markdown("---")
        st.markdown("### 💾 Download Report")
        docx_buffer = create_docx(summary, red_flags, simplified_chunks)
        st.download_button(
            label="📥 Download Simplified Report (.docx)",
            data=docx_buffer,
            file_name="simplified_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

elif uploaded_file and not api_key:
    st.warning("⚠️ Please enter your Groq API key in the sidebar to proceed.")
elif not uploaded_file:
    st.info("👆 Upload a PDF to get started.")